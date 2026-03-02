from datetime import datetime

from flask import Flask, render_template, request, redirect, url_for, send_file
from flask_sqlalchemy import SQLAlchemy

import io
import pandas as pd
from docx import Document


app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///bj_editor_stats.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)


# --- 常量與成員定義 ---

MEMBERS = [
    "解方",
    "顏笑",
    "李鴻健",
    "王紹薇",
    "汪峰",
    "魏紹坤",
    "柳佳妮",
    "趙德民",
    "孫舒陽",
    "王傲多",
]

SPECIAL_MEMBERS = {"顏笑", "趙德民"}
LEADER_NAME = "汪峰"


def classify_article(line: str) -> str:
    """根據稿件命名規則分類稿件類型，未識別返回 '其他'。"""
    text = line.upper()
    if "收片" in text:
        return "收片"
    if "SB+LVO" in text:
        return "SB+LVO"
    if "SB+ONLY" in text or "SBONLY" in text:
        return "SB+ONLY"
    if "SOT" in text:
        return "SOT"
    if "LVO" in text:
        return "LVO"
    # 乾稿/干+圖，名稱中常見「干」「干+圖」等，這裡做簡單匹配
    if "干+图" in text or "干+圖" in text:
        return "干+圖"
    if "干稿" in text or "干" in text:
        return "干稿"
    return "其他"


# --- 數據模型 ---


class DailyTask(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    member_name = db.Column(db.String(50), nullable=False)
    # 原始多行文字
    raw_text = db.Column(db.Text, nullable=False)
    # 完成日期（僅日期）
    task_date = db.Column(db.Date, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class InterpretingTask(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    member_name = db.Column(db.String(50), nullable=False)
    task_type = db.Column(db.String(20), nullable=False)  # 同傳 / 解說 / 同傳+解說
    time_range = db.Column(db.String(100), nullable=False)  # 文本形式保存，如 260301 09:00-11:30
    content = db.Column(db.Text, nullable=False)
    task_date = db.Column(db.Date, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


with app.app_context():
    db.create_all()


# --- 輔助函數 ---


def parse_daily_lines(raw_text: str):
    """將多行稿件文字拆分為清洗後的行列表。"""
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
    return lines


def aggregate_daily_statistics(tasks):
    """對給定的 DailyTask 集合做統計，返回 pandas.DataFrame。"""
    records = []
    for t in tasks:
        for line in parse_daily_lines(t.raw_text):
            category = classify_article(line)
            # 收片N 需要解析數量
            if category == "收片":
                # 例如 "收片2" 或 "收片 2"
                cnt = 1
                digits = "".join(ch for ch in line if ch.isdigit())
                if digits:
                    try:
                        cnt = int(digits)
                    except ValueError:
                        cnt = 1
                records.append(
                    {
                        "成員": t.member_name,
                        "日期": t.task_date.strftime("%Y-%m-%d"),
                        "類型": "收片",
                        "數量": cnt,
                        "詳情": line,
                    }
                )
            else:
                records.append(
                    {
                        "成員": t.member_name,
                        "日期": t.task_date.strftime("%Y-%m-%d"),
                        "類型": category,
                        "數量": 1,
                        "詳情": line,
                    }
                )

    if not records:
        return pd.DataFrame(columns=["成員", "日期", "類型", "數量", "詳情"])
    return pd.DataFrame(records)


def aggregate_interpreting_statistics(tasks):
    """對同傳/解說任務做統計，返回 DataFrame。"""
    records = []
    for t in tasks:
        records.append(
            {
                "成員": t.member_name,
                "日期": t.task_date.strftime("%Y-%m-%d"),
                "任務類別": t.task_type,
                "時間段": t.time_range,
                "內容": t.content,
            }
        )
    if not records:
        return pd.DataFrame(columns=["成員", "日期", "任務類別", "時間段", "內容"])
    return pd.DataFrame(records)


def is_leader(name: str) -> bool:
    return name == LEADER_NAME


def is_special_member(name: str) -> bool:
    return name in SPECIAL_MEMBERS


# --- 路由 ---


@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        name = request.form.get("member_name")
        if name not in MEMBERS:
            return render_template("login.html", members=MEMBERS, error="請選擇名單內成員。")
        return redirect(url_for("dashboard", member_name=name))
    return render_template("login.html", members=MEMBERS)


@app.route("/dashboard/<member_name>")
def dashboard(member_name):
    if member_name not in MEMBERS:
        return redirect(url_for("login"))
    return render_template(
        "dashboard.html",
        member_name=member_name,
        is_leader=is_leader(member_name),
        is_special=is_special_member(member_name),
    )


@app.route("/daily/<member_name>", methods=["GET", "POST"])
def daily_tasks(member_name):
    if member_name not in MEMBERS:
        return redirect(url_for("login"))

    if request.method == "POST":
        raw_text = request.form.get("raw_text", "").strip()
        date_str = request.form.get("task_date")
        if not raw_text or not date_str:
            error = "請填寫稿件內容並選擇日期。"
            tasks = (
                DailyTask.query.filter_by(member_name=member_name)
                .order_by(DailyTask.task_date.desc(), DailyTask.created_at.desc())
                .all()
            )
            return render_template(
                "daily_tasks.html",
                member_name=member_name,
                tasks=tasks,
                error=error,
            )
        task_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        task = DailyTask(member_name=member_name, raw_text=raw_text, task_date=task_date)
        db.session.add(task)
        db.session.commit()
        return redirect(url_for("daily_tasks", member_name=member_name))

    tasks = (
        DailyTask.query.filter_by(member_name=member_name)
        .order_by(DailyTask.task_date.desc(), DailyTask.created_at.desc())
        .all()
    )
    return render_template("daily_tasks.html", member_name=member_name, tasks=tasks)


@app.route("/daily/<member_name>/delete/<int:task_id>", methods=["POST"])
def delete_daily_task(member_name, task_id):
    task = DailyTask.query.get_or_404(task_id)
    if task.member_name != member_name:
        return redirect(url_for("daily_tasks", member_name=member_name))
    db.session.delete(task)
    db.session.commit()
    return redirect(url_for("daily_tasks", member_name=member_name))


@app.route("/interpreting/<member_name>", methods=["GET", "POST"])
def interpreting_tasks(member_name):
    if member_name not in MEMBERS or not is_special_member(member_name):
        return redirect(url_for("dashboard", member_name=member_name))

    if request.method == "POST":
        task_type = request.form.get("task_type", "").strip()
        time_range = request.form.get("time_range", "").strip()
        content = request.form.get("content", "").strip()
        date_str = request.form.get("task_date")
        if not (task_type and time_range and content and date_str):
            error = "請完整填寫任務類別、時間段、內容與日期。"
            tasks = (
                InterpretingTask.query.filter_by(member_name=member_name)
                .order_by(InterpretingTask.task_date.desc(), InterpretingTask.created_at.desc())
                .all()
            )
            return render_template(
                "interpreting_tasks.html",
                member_name=member_name,
                tasks=tasks,
                error=error,
            )
        task_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        task = InterpretingTask(
            member_name=member_name,
            task_type=task_type,
            time_range=time_range,
            content=content,
            task_date=task_date,
        )
        db.session.add(task)
        db.session.commit()
        return redirect(url_for("interpreting_tasks", member_name=member_name))

    tasks = (
        InterpretingTask.query.filter_by(member_name=member_name)
        .order_by(InterpretingTask.task_date.desc(), InterpretingTask.created_at.desc())
        .all()
    )
    return render_template("interpreting_tasks.html", member_name=member_name, tasks=tasks)


@app.route("/interpreting/<member_name>/delete/<int:task_id>", methods=["POST"])
def delete_interpreting_task(member_name, task_id):
    task = InterpretingTask.query.get_or_404(task_id)
    if task.member_name != member_name:
        return redirect(url_for("interpreting_tasks", member_name=member_name))
    db.session.delete(task)
    db.session.commit()
    return redirect(url_for("interpreting_tasks", member_name=member_name))


@app.route("/public")
def public_board():
    # 篩選條件
    member = request.args.get("member")
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")

    daily_query = DailyTask.query
    interp_query = InterpretingTask.query

    if member:
        daily_query = daily_query.filter_by(member_name=member)
        interp_query = interp_query.filter_by(member_name=member)

    if start_date:
        sd = datetime.strptime(start_date, "%Y-%m-%d").date()
        daily_query = daily_query.filter(DailyTask.task_date >= sd)
        interp_query = interp_query.filter(InterpretingTask.task_date >= sd)
    if end_date:
        ed = datetime.strptime(end_date, "%Y-%m-%d").date()
        daily_query = daily_query.filter(DailyTask.task_date <= ed)
        interp_query = interp_query.filter(InterpretingTask.task_date <= ed)

    daily_tasks = daily_query.order_by(DailyTask.task_date.desc(), DailyTask.created_at.desc()).all()
    interp_tasks = interp_query.order_by(
        InterpretingTask.task_date.desc(), InterpretingTask.created_at.desc()
    ).all()

    return render_template(
        "public_board.html",
        members=MEMBERS,
        daily_tasks=daily_tasks,
        interp_tasks=interp_tasks,
        selected_member=member or "",
        start_date=start_date or "",
        end_date=end_date or "",
    )


@app.route("/stats")
def stats():
    # 僅頁面權限提示在前端處理，後端不做限制（工具為內部公開使用）
    member = request.args.get("member")
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")

    daily_query = DailyTask.query
    interp_query = InterpretingTask.query
    if member:
        daily_query = daily_query.filter_by(member_name=member)
        interp_query = interp_query.filter_by(member_name=member)
    if start_date:
        sd = datetime.strptime(start_date, "%Y-%m-%d").date()
        daily_query = daily_query.filter(DailyTask.task_date >= sd)
        interp_query = interp_query.filter(InterpretingTask.task_date >= sd)
    if end_date:
        ed = datetime.strptime(end_date, "%Y-%m-%d").date()
        daily_query = daily_query.filter(DailyTask.task_date <= ed)
        interp_query = interp_query.filter(InterpretingTask.task_date <= ed)

    daily_tasks = daily_query.all()
    interp_tasks = interp_query.all()

    df_daily = aggregate_daily_statistics(daily_tasks)
    df_interp = aggregate_interpreting_statistics(interp_tasks)

    # 個人統計
    personal_stats = {}
    if not df_daily.empty:
        grp = df_daily.groupby(["成員", "類型"])["數量"].sum().reset_index()
        for _, row in grp.iterrows():
            name = row["成員"]
            personal_stats.setdefault(name, {})
            personal_stats[name][row["類型"]] = int(row["數量"])

    # 收片總數與稿件總數
    dept_totals = {}
    if not df_daily.empty:
        dept_grp = df_daily.groupby("類型")["數量"].sum().reset_index()
        for _, row in dept_grp.iterrows():
            dept_totals[row["類型"]] = int(row["數量"])

    interp_personal = {}
    if not df_interp.empty:
        grp_i = df_interp.groupby(["成員", "任務類別"]).size().reset_index(name="次數")
        for _, row in grp_i.iterrows():
            name = row["成員"]
            interp_personal.setdefault(name, {})
            interp_personal[name][row["任務類別"]] = int(row["次數"])

    interp_total = int(df_interp.shape[0]) if not df_interp.empty else 0

    return render_template(
        "stats.html",
        members=MEMBERS,
        personal_stats=personal_stats,
        dept_totals=dept_totals,
        interp_personal=interp_personal,
        interp_total=interp_total,
        start_date=start_date or "",
        end_date=end_date or "",
        selected_member=member or "",
        has_daily=not df_daily.empty,
        has_interp=not df_interp.empty,
    )


def build_export_frames(start_date=None, end_date=None):
    daily_query = DailyTask.query
    interp_query = InterpretingTask.query
    if start_date:
        sd = datetime.strptime(start_date, "%Y-%m-%d").date()
        daily_query = daily_query.filter(DailyTask.task_date >= sd)
        interp_query = interp_query.filter(InterpretingTask.task_date >= sd)
    if end_date:
        ed = datetime.strptime(end_date, "%Y-%m-%d").date()
        daily_query = daily_query.filter(DailyTask.task_date <= ed)
        interp_query = interp_query.filter(InterpretingTask.task_date <= ed)
    daily_tasks = daily_query.all()
    interp_tasks = interp_query.all()
    df_daily = aggregate_daily_statistics(daily_tasks)
    df_interp = aggregate_interpreting_statistics(interp_tasks)
    return df_daily, df_interp


@app.route("/export/<fmt>")
def export(fmt: str):
    # 所有人可下載，使用日期篩選
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    df_daily, df_interp = build_export_frames(start_date, end_date)

    today_tag = datetime.now().strftime("%y%m%d")
    base_name = f"北京編譯中心稿件統計_{today_tag}_{fmt}"

    if fmt == "txt":
        output = io.StringIO()
        output.write("【日常稿件與收片】\n")
        if df_daily.empty:
            output.write("無數據\n")
        else:
            for _, row in df_daily.iterrows():
                output.write(
                    f"{row['日期']} {row['成員']} [{row['類型']}] x{row['數量']} {row['詳情']}\n"
                )
        output.write("\n【同傳/解說任務】\n")
        if df_interp.empty:
            output.write("無數據\n")
        else:
            for _, row in df_interp.iterrows():
                output.write(
                    f"{row['日期']} {row['成員']} [{row['任務類別']}] {row['時間段']} {row['內容']}\n"
                )
        mem = io.BytesIO(output.getvalue().encode("utf-8"))
        mem.seek(0)
        return send_file(
            mem,
            as_attachment=True,
            download_name=f"{base_name}.txt",
            mimetype="text/plain; charset=utf-8",
        )

    if fmt in ("xlsx", "excel"):
        mem = io.BytesIO()
        with pd.ExcelWriter(mem, engine="openpyxl") as writer:
            df_daily.to_excel(writer, index=False, sheet_name="日常稿件與收片")
            df_interp.to_excel(writer, index=False, sheet_name="同傳與解說")
        mem.seek(0)
        return send_file(
            mem,
            as_attachment=True,
            download_name=f"{base_name}.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    if fmt in ("docx", "word"):
        doc = Document()
        doc.add_heading("北京編譯中心稿件統計", level=1)
        doc.add_paragraph()
        doc.add_heading("日常稿件與收片", level=2)
        if df_daily.empty:
            doc.add_paragraph("無數據")
        else:
            for _, row in df_daily.iterrows():
                doc.add_paragraph(
                    f"{row['日期']} {row['成員']} [{row['類型']}] x{row['數量']} {row['詳情']}"
                )
        doc.add_paragraph()
        doc.add_heading("同傳/解說任務", level=2)
        if df_interp.empty:
            doc.add_paragraph("無數據")
        else:
            for _, row in df_interp.iterrows():
                doc.add_paragraph(
                    f"{row['日期']} {row['成員']} [{row['任務類別']}] {row['時間段']} {row['內容']}"
                )
        mem = io.BytesIO()
        doc.save(mem)
        mem.seek(0)
        return send_file(
            mem,
            as_attachment=True,
            download_name=f"{base_name}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # 未知格式
    return redirect(url_for("stats"))


if __name__ == "__main__":
    app.run(debug=True)

